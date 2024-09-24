import tkinter as tk
from tkinter import filedialog, messagebox
import speech_recognition as sr
from moviepy.editor import VideoFileClip, AudioFileClip
import timeit
import tempfile
import os
from docx import Document
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline

# Load pre-trained model and tokenizer
model_name = "oliverguhr/fullstop-punctuation-multilang-large"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForTokenClassification.from_pretrained(model_name)

# Define pipeline for punctuation restoration
punctuation_pipeline = pipeline("ner", model=model, tokenizer=tokenizer, grouped_entities=True)

def add_punctuation(text):
    result = punctuation_pipeline(text)
    punctuated_text = ""
    for res in result:
        word = res.get('word', '')
        entity_group = res.get('entity_group', 'O')
        punctuated_text += word
        if entity_group == '0':
            punctuated_text += " "
        else:
            punctuation = entity_group if entity_group else ''
            punctuated_text += punctuation + " "
    return punctuated_text.strip()

def video_to_audio(video_file_path, audio_file_path):
    try:
        video_clip = VideoFileClip(video_file_path)
        audio_clip = video_clip.audio
        if audio_clip is not None:
            audio_clip.write_audiofile(audio_file_path, codec='pcm_s16le')
            video_clip.close()
            audio_clip.close()
        else:
            raise ValueError("The selected video file does not contain an audio track.")
    except Exception as e:
        raise e

def convert_audio_to_text(audio_path):
    r = sr.Recognizer()
    if audio_path.endswith(".mp3"):
        audio_clip = AudioFileClip(audio_path)
        audio_path_wav = audio_path.replace(".mp3", ".wav")
        audio_clip.write_audiofile(audio_path_wav, codec='pcm_s16le')
        audio_clip.close()
        audio_path = audio_path_wav

    with sr.AudioFile(audio_path) as source:
        audio = r.record(source)

    try:
        start_time = timeit.default_timer()
        text = r.recognize_google(audio)
        elapsed_time = timeit.default_timer() - start_time
        punctuated_text = add_punctuation(text)
        return punctuated_text, elapsed_time
    except Exception as e:
        return str(e), None

def open_file_dialog():
    global video_path
    video_path = filedialog.askopenfilename(title="Select a video file", filetypes=[("Video Files", "*.mp4 *.avi *.mov")])
    if video_path:
        file_label.config(text=f"Selected file: {os.path.basename(video_path)}")
    else:
        file_label.config(text="No file selected")

def convert_file():
    if video_path:
        try:
            temp_audio_fd, temp_audio_path = tempfile.mkstemp(suffix=".wav")
            os.close(temp_audio_fd)

            video_to_audio(video_path, temp_audio_path)
            text, elapsed_time = convert_audio_to_text(temp_audio_path)

            text_box.delete(1.0, tk.END)
            text_box.insert(tk.END, text)
            if elapsed_time is not None:
                time_label.config(text=f"Conversion Time: {elapsed_time:.2f} seconds")
            else:
                time_label.config(text="Conversion Time: Error occurred")

            os.remove(temp_audio_path)
        except Exception as e:
            messagebox.showerror("Error", str(e))
    else:
        messagebox.showwarning("Warning", "Please select a video file first")

def save_to_word():
    text = text_box.get(1.0, tk.END).strip()
    if not text:
        messagebox.showwarning("Warning", "No text to save!")
        return
    
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if save_path:
        try:
            document = Document()
            document.add_heading("Converted Text", level=1)
            document.add_paragraph(text)
            document.save(save_path)
            messagebox.showinfo("Success", "Text saved to Word document successfully!")
        except Exception as e:
            messagebox.showerror("Error", str(e))

root = tk.Tk()
root.title("Video to Text Converter")
root.geometry("800x600")
root.configure(bg="#f0f0f0")

# Create frames
header_frame = tk.Frame(root, bg="#003366", pady=10)
header_frame.pack(fill=tk.X)

content_frame = tk.Frame(root, bg="#f0f0f0", pady=20)
content_frame.pack(fill=tk.BOTH, expand=True)

footer_frame = tk.Frame(root, bg="#003366", pady=10)
footer_frame.pack(fill=tk.X, side=tk.BOTTOM)

# Header widgets
title_label = tk.Label(header_frame, text="Video to Text Converter", width=20,
                       height=2, borderwidth=4, relief="ridge", font=("Helvetica", 20), bg="#003366", fg="white")
title_label.pack(pady=10)

# Content widgets
file_label = tk.Label(content_frame, text="No file selected", font=("Helvetica", 12), bg="#f0f0f0", fg="#003366")
file_label.pack(pady=10)

open_button = tk.Button(content_frame, text="Open Video File", command=open_file_dialog, bg="#4CAF50", fg="white")
open_button.pack(pady=10)

convert_button = tk.Button(content_frame, text="Convert", command=convert_file, bg="#008CBA", fg="white")
convert_button.pack(pady=20)

save_button = tk.Button(content_frame, text="Save to Word", command=save_to_word, bg="#FF5733", fg="white")
save_button.pack(pady=10)

text_box = tk.Text(content_frame, wrap=tk.WORD, width=80, height=20, bg="#ffffff", fg="#000000")
text_box.pack(pady=20)

time_label = tk.Label(content_frame, text="Conversion Time: N/A", font=("Helvetica", 12), bg="#f0f0f0", fg="#003366")
time_label.pack(pady=10)

root.mainloop()
